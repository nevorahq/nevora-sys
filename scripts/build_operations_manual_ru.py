from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Nevora_Business_OS_Operations_Manual_ru.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 33, 45)
MUTED = RGBColor(89, 99, 110)
LIGHT_FILL = "E8EEF5"
PALE_FILL = "F4F6F9"
GRID = "D9E2EC"


@dataclass
class Workflow:
    name: str
    purpose: str
    business_value: str
    user_goal: str
    expected_result: str
    modules: str
    related_entities: str
    preconditions: list[str]
    journey: list[str]
    system_behaviour: list[str]
    ui_changes: list[str]
    business_rules: list[str]
    edge_cases: list[str]


def set_run_font(run, *, name: str = "Calibri", size: float | None = None,
                 color: RGBColor | None = None, bold: bool | None = None,
                 italic: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6,
                          line_spacing: float = 1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_number(paragraph):
    paragraph.add_run("Страница ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    set_paragraph_spacing(p, before={1: 18, 2: 14, 3: 10}.get(level, 8),
                          after={1: 10, 2: 7, 3: 5}.get(level, 4), line_spacing=1.25)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4)
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_numbers(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4)
        r = p.add_run(item)
        set_run_font(r, color=INK)


def add_label_table(doc: Document, rows: list[tuple[str, str]], widths: list[float] = [1.7, 4.8]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        set_cell_shading(cells[0], PALE_FILL)
        cells[0].paragraphs[0].add_run(label)
        cells[1].paragraphs[0].add_run(value)
        for cell in cells:
            for p in cell.paragraphs:
                set_paragraph_spacing(p, after=0, line_spacing=1.15)
                for run in p.runs:
                    set_run_font(run, size=10.5, color=INK)
            if cell is cells[0]:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, PALE_FILL)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line_spacing=1.15)
    r = p.add_run(title + ": ")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        s.header_distance = Inches(0.492)
        s.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 26, INK, 0, 4),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def setup_header_footer(doc: Document):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header.paragraphs[0]
    header.text = "Nevora Business OS | Operations Manual"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=0, line_spacing=1.0)
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)
    set_paragraph_spacing(footer, after=0, line_spacing=1.0)
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)


def add_cover(doc: Document):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nevora Business OS")
    set_run_font(r, size=30, color=INK, bold=True)
    set_paragraph_spacing(p, after=4, line_spacing=1.15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Операционное руководство пользователя")
    set_run_font(r, size=18, color=DARK_BLUE, bold=True)
    set_paragraph_spacing(p, after=10, line_spacing=1.15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Product Documentation | User Guide | QA Manual | Support Wiki")
    set_run_font(r, size=11, color=MUTED, bold=True)
    set_paragraph_spacing(p, after=22, line_spacing=1.15)

    add_callout(
        doc,
        "Назначение",
        "Документ объясняет, как пользователь работает в Nevora, какие модули участвуют в каждом сценарии и что система делает внутри: проверки прав, записи в базу, события, уведомления, AI-предложения, фоновые задания и защиту от повторных действий.",
    )
    add_label_table(doc, [
        ("Версия", "1.0"),
        ("Дата подготовки", "12 июля 2026"),
        ("Статус продукта", "MVP / private-beta stabilization"),
        ("Язык", "Русский"),
        ("Источник", "Репозиторий nevora-sys, product contracts, module status, workflow plan"),
    ])
    doc.add_page_break()


MODULE_ROWS = [
    ("Action Center", "Главный ежедневный экран: что требует внимания сегодня.", "/dashboard", "MVP Ready"),
    ("Dashboard Overview", "Сводные метрики по активным областям бизнеса.", "/dashboard/overview", "MVP Ready"),
    ("Capture Inbox", "Быстрый ввод идей, обязательств и заметок с AI-предложениями.", "/dashboard/inbox", "MVP Ready"),
    ("Tasks", "Обычные задачи, проекты, исполнители, сроки, статусы.", "/dashboard/tasks", "MVP Ready"),
    ("Financial Tasks", "Платёжные обязательства, которые можно явно отметить как оплаченные.", "/dashboard/tasks/financial", "MVP Ready"),
    ("Money", "Счета, доходы, расходы, переводы, категории, плановые операции.", "/dashboard/money", "MVP Ready"),
    ("Documents", "Документы, вложения, извлечение данных, черновики финансовых действий.", "/dashboard/documents", "MVP Ready"),
    ("Subscriptions", "Подписки, циклы платежей, напоминания о списаниях.", "/dashboard/subscriptions", "MVP Ready"),
    ("Notifications", "Уведомления, push, напоминания, read-state.", "Bell + settings", "MVP Ready"),
    ("Settings", "Профиль, workspace, участники, биллинг, уведомления, developer access.", "/dashboard/settings", "MVP Ready / Partial"),
    ("AI", "Инсайты, рекомендации, summary, extraction, inbox intent detection.", "/dashboard/ai", "Partial"),
    ("Analytics", "Метрики, snapshots, reports, activity timeline.", "/dashboard/analytics", "Partial"),
    ("Relations", "Связи между задачами, документами, транзакциями и подписками.", "Inline", "In Progress"),
    ("Automation", "Обработчики доменных событий и фоновые repair jobs.", "No direct page", "Foundation"),
    ("CRM", "Клиенты, сделки, pipeline.", "/dashboard/crm", "Paused / hard-gated"),
    ("Booking", "Публичная запись на услуги.", "/dashboard/booking, /booking/*", "Paused / hard-gated"),
]


WORKFLOWS = [
    Workflow(
        name="Регистрация, вход и создание рабочей организации",
        purpose="Дать пользователю доступ к защищённой рабочей среде Nevora и связать все последующие действия с конкретной организацией.",
        business_value="Надёжная tenant-изоляция позволяет одной SaaS-платформе обслуживать несколько компаний без смешивания данных.",
        user_goal="Создать аккаунт, войти, пройти onboarding и попасть в свой рабочий dashboard.",
        expected_result="Пользователь авторизован, организация и workspace созданы или выбраны, маршруты dashboard доступны.",
        modules="Auth, Organizations, Workspaces, Settings, Billing Access",
        related_entities="profiles, organizations, workspaces, memberships, invites",
        preconditions=[
            "Пользователь открывает публичный маршрут /register, /login или invite-ссылку.",
            "Для защищённых маршрутов требуется валидная Supabase-сессия.",
            "Если у пользователя ещё нет организации, он должен пройти /onboarding.",
            "Если пользователь пришёл по приглашению, token должен быть действительным и не использованным.",
        ],
        journey=[
            "Экран регистрации: пользователь вводит email и пароль, нажимает кнопку создания аккаунта. Система валидирует формат email, обязательные поля и ответ Supabase Auth.",
            "Экран входа: пользователь вводит учётные данные и нажимает вход. После успешной сессии proxy пропускает защищённые маршруты.",
            "Экран onboarding: пользователь вводит название организации и подтверждает создание. Система создаёт organization, workspace и membership через серверный контекст.",
            "Invite flow: пользователь открывает /invite/[token], просматривает информацию о приглашении и принимает его. Система связывает пользователя с организацией без доверия к client-provided organization_id.",
            "После успешного входа пользователь попадает на /dashboard, где главным экраном является Action Center.",
        ],
        system_behaviour=[
            "proxy.ts обновляет сессию Supabase, пропускает публичные и machine routes, защищённые маршруты без сессии отправляет на /login.",
            "requireOrg() на сервере определяет active membership, organization, workspace и role-derived permissions.",
            "RLS в PostgreSQL остаётся финальной границей: строки доступны только членам соответствующей организации.",
            "SECURITY DEFINER RPC для provisioning и invite flow имеет явные GRANT EXECUTE и фиксированный search_path.",
            "Billing access state может ограничить write/execute действия, если trial истёк или организация стала read-only.",
        ],
        ui_changes=[
            "Публичные страницы сменяются onboarding или dashboard в зависимости от состояния пользователя.",
            "При ошибке входа пользователь остаётся на форме и видит сообщение об ошибке.",
            "После создания организации отображается рабочая навигация dashboard.",
            "Если доступ запрещён, пользователь видит controlled refusal, а не чужие данные.",
        ],
        business_rules=[
            "organization_id и workspace_id никогда не берутся из формы как источник истины.",
            "Owner/admin/member права выводятся из membership и permission set.",
            "Machine routes не используют user session, но обязаны проверять собственные секреты или подписи.",
            "Один пользователь может работать только с данными активной организации.",
        ],
        edge_cases=[
            "Истёкшая сессия: защищённый маршрут отправляет пользователя на /login.",
            "Пользователь без организации: dashboard недоступен до onboarding.",
            "Недействительный invite token: приглашение не принимается, membership не создаётся.",
            "Попытка подменить organization_id в payload: сервер игнорирует клиентский tenant context, RLS дополнительно блокирует доступ.",
            "Переключение организации: последующие queries и actions должны пересчитать active context.",
        ],
    ),
    Workflow(
        name="Ежедневная работа через Action Center",
        purpose="Собрать все сигналы из модулей в один список действий: что проверить, оплатить, завершить, уточнить или отложить.",
        business_value="SMB-пользователь начинает день с приоритетного списка вместо обхода всех разделов вручную.",
        user_goal="Понять, что требует внимания сегодня, и закрыть элементы через confirm, resolve, dismiss, snooze или assign.",
        expected_result="Нужные action_items обработаны, бизнес-состояние изменено только после явного действия пользователя.",
        modules="Action Center, Tasks, Money, Documents, Subscriptions, Capture Inbox, Notifications",
        related_entities="action_items, action_item_events, todos, money_transactions, documents, subscriptions, notifications",
        preconditions=[
            "Пользователь вошёл в организацию и открывает /dashboard.",
            "В системе есть задачи, подписки, документы, финансовые черновики или AI-предложения.",
            "Пользователь имеет права читать соответствующие элементы и выполнять доступные действия.",
        ],
        journey=[
            "Экран /dashboard: пользователь видит четыре секции - Needs your review, Money attention, Next actions, Recently updated.",
            "Пользователь открывает карточку или detail drawer, читает описание, источник, deadline, приоритет и связанные сущности.",
            "Если элемент требует решения, пользователь выбирает confirm/execute, reject/dismiss, snooze, resolve или assign. Система показывает состояние обработки.",
            "После подтверждения система выполняет действие через модульный сервис, обновляет action_item и возвращает пользователя к списку.",
            "Пользователь проверяет summary strip: счётчики overdue, due today, upcoming и recently resolved должны измениться согласно бизнес-состоянию.",
        ],
        system_behaviour=[
            "syncActionItems() может best-effort синхронизировать элементы при загрузке страницы; операция идемпотентна и не должна ломать рендер.",
            "Action items дедуплицируются по organization, type, source_type и source_id.",
            "Переходы статусов пишут action_items.status, создают action item events и доменные события.",
            "Money attention определяется по типу payment/renewal или source/primary entity transaction/subscription.",
            "Read-state уведомления не влияет на action_items: read is not resolved.",
        ],
        ui_changes=[
            "Карточка может перейти из open в resolved, dismissed или snoozed.",
            "Snoozed элемент исчезает из активной секции до snoozed_until.",
            "Resolved элемент появляется в Recently updated.",
            "Ошибки permission или validation отображаются без частичного изменения бизнес-данных.",
        ],
        business_rules=[
            "Action Center не является ledger: денежная транзакция появляется только через модуль Money или утверждённый idempotent workflow.",
            "Dismiss означает 'не актуально', а resolve означает 'обработано'.",
            "Секция Money attention имеет приоритет для финансовых объектов, даже если базовый тип элемента относится к review.",
            "Старый маршрут /dashboard/actions должен перенаправлять на /dashboard для совместимости.",
        ],
        edge_cases=[
            "Double click на confirm: модульный сервис или RPC должен быть идемпотентным, особенно для денег.",
            "Сетевой сбой после подтверждения: пользователь обновляет страницу, а action item должен показать фактическое состояние.",
            "Связанная сущность удалена: карточка не должна падать, связь должна быть пропущена или показана как недоступная.",
            "Пользователь читает notification: action item остаётся активным.",
            "Concurrent update другим участником: текущий пользователь должен увидеть обновлённый статус при следующем refresh/revalidate.",
        ],
    ),
    Workflow(
        name="Capture Inbox: быстрый ввод и AI-предложения",
        purpose="Позволить пользователю быстро записать текст, идею, обязательство или напоминание, а системе предложить безопасное действие.",
        business_value="Снижает потерю операционных сигналов и превращает сырой ввод в reviewable work без автономных финансовых записей.",
        user_goal="Добавить запись в Inbox, получить предложение, принять, отредактировать или отклонить его.",
        expected_result="После подтверждения создаётся задача, финансовая задача, связь или action item; posted money transaction не создаётся.",
        modules="Capture Inbox / Planner, AI, Tasks, Action Center, Relations",
        related_entities="planner_entries, planner_suggestions, todos, entity_links, action_items",
        preconditions=[
            "Пользователь находится на /dashboard/inbox.",
            "У пользователя есть права planner.entry.create и права на целевое действие.",
            "Ввод содержит raw_text или связанную source entity.",
        ],
        journey=[
            "Экран Inbox: пользователь нажимает создание capture-записи, вводит текст и сохраняет.",
            "Система переводит запись в processing, определяет intent и confidence.",
            "Пользователь видит pending suggestion: создать задачу, финансовую задачу, напоминание, связь или action item.",
            "Пользователь открывает suggestion, проверяет title, description и proposed payload.",
            "Пользователь нажимает accept, edit или reject. Accept создаёт сущность только через существующий сервис целевого модуля.",
        ],
        system_behaviour=[
            "planner_entries хранит raw input и status captured -> processing -> suggested/failed.",
            "planner_suggestions хранит только reviewable proposals с allow-list типов.",
            "Accept revalidates payload per suggestion type; proposed_payload никогда не mass-assigned.",
            "Финансовые предложения маршрутизируются в createFinancialTask и никогда не пишут money_transactions напрямую.",
            "На suggestion creation или failure создаётся Action Center item: ai_suggestion или missing_information.",
        ],
        ui_changes=[
            "Новая запись появляется в Inbox со статусом обработки.",
            "Pending suggestion показывает confidence, описание и доступные действия.",
            "Accepted suggestion показывает созданную сущность или ссылку на неё.",
            "Rejected suggestion сохраняет причину и не исчезает бесследно из audit trail.",
        ],
        business_rules=[
            "AI suggests, user confirms, module service executes.",
            "Разрешённые suggestion types не включают transaction/expense/income producer.",
            "Accept помечает suggestion accepted только после успешного создания целевой сущности.",
            "Delete заменяется архивированием/статусом, чтобы сохранить историю решений.",
        ],
        edge_cases=[
            "AI не понял intent: entry получает failed или missing-information action item.",
            "Invalid AI payload: fallback или отказ без создания business entity.",
            "Пользователь принимает suggestion дважды: compare-and-swap status transition не должен создать дубль.",
            "Недостаточно прав на целевую сущность: suggestion остаётся pending или получает controlled failure.",
            "Пользователь переключил организацию: accept должен выполняться в новом server-derived context или отказать.",
        ],
    ),
    Workflow(
        name="Обычные задачи и проекты",
        purpose="Управлять повседневной работой: создавать задачи, назначать исполнителей, менять статус, сроки и привязку к проектам.",
        business_value="Команда получает единый список работы, Action Center видит deadlines и overdue, а история действий остаётся отслеживаемой.",
        user_goal="Создать задачу, довести её от todo до done, при необходимости привязать к проекту и участникам.",
        expected_result="Задача обновлена, статусы и счётчики пересчитаны, связанные action items отражают новое состояние.",
        modules="Tasks, Projects, Members, Action Center, Notifications",
        related_entities="todos, task_projects, task_assignees, task_due_date_changes, action_items, notifications",
        preconditions=[
            "Пользователь открывает /dashboard/tasks или страницу проекта.",
            "У пользователя есть data.write для создания/изменения задач.",
            "Для назначения исполнителя участник должен принадлежать текущей организации.",
        ],
        journey=[
            "Экран Tasks: пользователь нажимает кнопку создания задачи, вводит title, description, priority, due date и сохраняет.",
            "Список задач обновляется; пользователь может открыть карточку или detail page /dashboard/tasks/[taskId].",
            "Пользователь меняет статус todo -> in_progress -> done через доступный контрол статуса.",
            "Пользователь добавляет исполнителя, меняет срок или привязывает задачу к проекту. Система валидирует значения и права.",
            "Если срок приближается или просрочен, Action Center и Notifications создают соответствующие элементы внимания.",
        ],
        system_behaviour=[
            "status является источником истины; is_completed остаётся backward-compatible зеркалом done.",
            "Due date changes классифицируются как set, extended, shortened, changed или removed.",
            "Task actions валидируют input через Zod и берут organization/workspace из requireOrg.",
            "Assignments, comments и activity feed привязаны к task и org context.",
            "Action Center генератор создаёт due_soon/overdue/follow_up_required элементы с дедупликацией.",
        ],
        ui_changes=[
            "Новая задача появляется в списке и summary cards.",
            "Статусный badge меняется на Not set, In progress или Closed.",
            "При изменении срока обновляется due label и возможный overdue state.",
            "Проектный экран обновляет progress bar и список задач.",
        ],
        business_rules=[
            "Завершение обычной задачи никогда не означает оплату.",
            "Task title ограничен 200 символами, description 2000, comment 5000.",
            "Активные статусы: todo и in_progress; completed: done.",
            "Сроки и reminders должны учитывать timezone пользователя/организации.",
        ],
        edge_cases=[
            "Удалённый исполнитель: задача не должна раскрыть чужие данные и должна безопасно обновить assignment.",
            "Browser refresh во время сохранения: повторный submit не должен ломать статус.",
            "Concurrent статус-апдейт: последний server action должен отразиться после revalidation.",
            "Срок удалён: будущие due reminders должны быть отменены.",
            "Пользователь без прав: action возвращает отказ без изменений.",
        ],
    ),
    Workflow(
        name="Финансовые задачи и Mark as paid",
        purpose="Отделить плановое обязательство от фактически проведённого платежа и создать posted transaction только после подтверждения.",
        business_value="Платформа помогает не забыть оплату, но не превращает AI, документ или задачу в бухгалтерский факт без человека.",
        user_goal="Проверить финансовую задачу и явно отметить её оплаченной, когда платёж действительно произошёл.",
        expected_result="Создана одна posted expense transaction, финансовая задача закрыта как paid, повторный клик не создаёт дубль.",
        modules="Tasks, Money, Documents, Subscriptions, Action Center",
        related_entities="todos, money_transactions, entity_links, action_items, domain_events",
        preconditions=[
            "У задачи task_context_type относится к payable context: subscription_payment, invoice_payment, tax_payment, domain_renewal или hosting_payment.",
            "financial_status = open.",
            "Сумма, валюта и источник обязательства доступны на сервере.",
        ],
        journey=[
            "Пользователь открывает Financial Tasks или Action Center item с payment_required.",
            "Пользователь раскрывает task panel и проверяет сумму, валюту, due date, источник и связанные документы.",
            "Пользователь нажимает Mark as paid. Система не принимает amount/currency как доверенный клиентский факт.",
            "Сервер вызывает защищённый workflow, создаёт posted transaction и связывает её с задачей.",
            "Пользователь видит paid/done state, а action item переходит в resolved.",
        ],
        system_behaviour=[
            "Для one-off financial tasks guard хранится в todos.financial_transaction_id.",
            "Если financial_transaction_id уже есть, повторный Mark as paid должен вернуть существующее состояние.",
            "Сервис создаёт money_transactions.status = posted только после explicit confirmation.",
            "Domain event и entity links создаются best-effort после primary write.",
            "Action Center и Notifications обновляются на основе бизнес-состояния, не read_at.",
        ],
        ui_changes=[
            "Кнопка Mark as paid становится недоступной или заменяется paid state.",
            "Финансовая задача исчезает из active obligation list или перемещается в completed/history.",
            "В Money появляется расход с соответствующей суммой.",
            "Если операция уже оплачена, UI должен показывать already paid, а не ошибку-дубль.",
        ],
        business_rules=[
            "Task completion alone is not payment.",
            "Planned obligation is not posted ledger entry.",
            "Only explicit confirmation or approved idempotent workflow may write posted money.",
            "Amount должен быть положительным; тип expense определяет знак в аналитике.",
        ],
        edge_cases=[
            "Double click: должен вернуться тот же transaction_id.",
            "Network retry после успешной оплаты: повторный запрос не создаёт второй расход.",
            "Финансовая задача уже dismissed/skipped: Mark as paid запрещён или требует восстановления.",
            "Связанный документ удалён: оплата может быть запрещена или выполнена без preview, но без падения UI.",
            "Недостаточно прав или read-only billing state: posted transaction не создаётся.",
        ],
    ),
    Workflow(
        name="Money: счета, транзакции, переводы и категории",
        purpose="Вести фактические доходы/расходы, плановые операции, переводы и категоризацию в рамках организации.",
        business_value="Даёт владельцу малого бизнеса обзор денег, предстоящих обязательств и категорий расходов без смешивания валют.",
        user_goal="Создать счёт, добавить доход/расход, провести перевод, проверить категории и прогнозы.",
        expected_result="Баланс, breakdown, recent transactions и Action Center обновлены согласно операции.",
        modules="Moneyflow, Action Center, AI suggestions, Documents",
        related_entities="money_accounts, money_transactions, money_categories, money_ai_suggestions, entity_links",
        preconditions=[
            "Пользователь находится на /dashboard/money.",
            "Для транзакции требуется активный счёт и валюта из EUR, USD, MDL, RUB.",
            "Для posted transaction пользователь должен явно нажать создание или подтверждение.",
        ],
        journey=[
            "Пользователь создаёт Money account: вводит name, type, initial balance, currency и сохраняет.",
            "Пользователь добавляет transaction: выбирает income или expense, account, amount, date, title, optional category/note.",
            "Для transfer пользователь выбирает from account, to account, amount и дату. Система создаёт нейтральный transfer row, не категорию income/expense.",
            "Пользователь открывает uncategorized transactions или category intelligence, принимает, редактирует или отклоняет suggestion.",
            "Пользователь проверяет summary cards, expense breakdown, recent transactions и planned transactions.",
        ],
        system_behaviour=[
            "transaction amount хранится положительным числом; type income/expense определяет знак для баланса.",
            "status posted входит в баланс, status planned остаётся прогнозом.",
            "transfer исключается из income/expense analytics и хранит from_account_id/to_account_id.",
            "Multi-currency summary показывает byCurrency и base summary через historical FX, если курс доступен.",
            "AI/category suggestions остаются pending до user review; user rules могут применяться напрямую.",
        ],
        ui_changes=[
            "После создания счёта он появляется в accounts list.",
            "После posted transaction меняются balance, monthly income/expenses и recent list.",
            "Planned transaction появляется в Upcoming/forecast, но не увеличивает ledger balance.",
            "Категоризация меняет badge uncategorized/suggested/confirmed.",
        ],
        business_rules=[
            "Нельзя суммировать разные валюты без FX layer; если курс отсутствует, base summary incomplete.",
            "Posting foreign-currency document draft на mismatched account блокируется.",
            "AI suggestion is not accounting fact.",
            "Account deactivation не должна удалять историю транзакций.",
        ],
        edge_cases=[
            "Счёт деактивирован до создания транзакции: форма должна отказать.",
            "Amount <= 0: validation error.",
            "Отсутствует FX rate: показывать per-currency breakdown и incomplete base, а не неверную сумму.",
            "Transfer between same account: validation error.",
            "Duplicate submit: user should not accidentally create duplicate posted rows; UI должен показывать loading/disabled state.",
        ],
    ),
    Workflow(
        name="Documents: загрузка, извлечение данных и подтверждение расхода",
        purpose="Хранить документы и превращать финансовые документы в проверяемые черновики действий.",
        business_value="Сокращает ручной ввод, но сохраняет принцип confirm-first для денежных операций.",
        user_goal="Загрузить чек/инвойс/подтверждение оплаты, проверить распознанные данные и подтвердить или отклонить транзакцию.",
        expected_result="Документ сохранён приватно, AI extraction создаёт review item, money transaction появляется только после подтверждения.",
        modules="Documents, Storage, AI extraction, Money, Action Center, Cron",
        related_entities="documents, document_attachments, extraction jobs, money_transactions, action_items, entity_links",
        preconditions=[
            "Пользователь открывает /dashboard/documents или /dashboard/documents/new.",
            "Файл имеет разрешённый тип: pdf, docx, png, jpg, jpeg, webp, heic, heif.",
            "Размер одного файла не больше 10 MB, максимум 5 файлов и 25 MB суммарно.",
            "Для extraction нужен Anthropic API или DOCUMENT_EXTRACTION_MOCK в локальной среде.",
        ],
        journey=[
            "Пользователь нажимает создание/загрузку документа, вводит title, optional description, doc_type и выбирает файлы.",
            "Система валидирует metadata и файлы, затем сохраняет attachment в приватное хранилище.",
            "Если doc_type финансовый: receipt, invoice или payment_confirmation, система ставит extraction job.",
            "Пользователь открывает document detail и видит статус извлечения, preview/attachment и review panel.",
            "Пользователь подтверждает, редактирует или отклоняет предложенную транзакцию. Только confirm создаёт posted или planned money row согласно workflow.",
        ],
        system_behaviour=[
            "Storage upload проходит feature gate и usage accounting.",
            "Extraction job claim state machine защищает от двойной обработки: pending -> processing -> done/failed.",
            "after() fast path и extraction-sweep могут конкурировать, но claim update должен дать работу только одному процессору.",
            "confirmDocumentTransactionAction валидирует account/currency/category и создаёт транзакцию только после решения пользователя.",
            "Document detail показывает reverse linked entities через UniversalRelationViewer.",
        ],
        ui_changes=[
            "Документ появляется в list/detail с draft/published/archived status.",
            "Extraction status меняется с pending/processing на done или failed.",
            "Action Center показывает draft_review/document_review/missing_information item.",
            "После confirm review panel должен показать resolved state и ссылку на money transaction.",
        ],
        business_rules=[
            "Document detection is not payment.",
            "Attaching a document is not posting money.",
            "Financial extraction accepts statement/unknown types in schema, but auto-extraction MVP targets receipt, invoice, payment_confirmation.",
            "Private bucket RLS must protect files by organization.",
        ],
        edge_cases=[
            "Файл слишком большой или тип запрещён: upload отклоняется до записи.",
            "Extraction stuck: cron extraction-sweep должен recover/retry.",
            "AI вернул неоднозначные данные: пользователь получает missing information или editable draft.",
            "Пользователь отклоняет draft: money transaction не создаётся.",
            "Документ уже удалён или archived: связанный action item должен безопасно обновиться.",
        ],
    ),
    Workflow(
        name="Subscriptions: подписки, циклы платежей и оплата",
        purpose="Отслеживать регулярные платежи и проводить расход только при явном Mark as paid.",
        business_value="Пользователь заранее видит списания, а система исключает двойную оплату через идемпотентный цикл.",
        user_goal="Создать подписку, видеть следующий платёж, получить payment task и отметить фактическую оплату.",
        expected_result="Создан planned payment cycle, затем ровно одна posted transaction после Mark as paid; подписка продвинута на следующий период.",
        modules="Subscriptions, Tasks, Money, Action Center, Cron, Relations",
        related_entities="subscriptions, subscription_payment_cycles, todos, money_transactions, entity_links, action_items",
        preconditions=[
            "Пользователь находится на /dashboard/subscriptions.",
            "Подписка имеет amount, currency, billing_cycle и next_billing_date.",
            "Для оплаты пользователь имеет права data.write и доступ к money workflow.",
        ],
        journey=[
            "Пользователь создаёт subscription: вводит название, сумму, валюту, billing cycle, дату следующего платежа и optional category/url/note.",
            "Система создаёт или переиспользует planned payment cycle для due date.",
            "Action Center показывает renewal/payment item или suggestion на создание payment task.",
            "Пользователь создаёт/подтверждает payment task, затем в нужный момент нажимает Mark as paid.",
            "Система создаёт posted expense, помечает cycle paid, закрывает task и рассчитывает следующий cycle.",
        ],
        system_behaviour=[
            "createSubscriptionPaymentCycle идемпотентен по billing_period_key и open cycle.",
            "mark_subscription_payment_paid RPC использует FOR UPDATE row lock, status guard и unique idempotency key.",
            "subscription-sweep ежедневно repair-only: создаёт missing cycles/suggestions/action items, но не платит.",
            "Billing period key и anchor-day сохраняют корректную дату при weekly/monthly/yearly циклах.",
            "entity_links связывают cycle/subscription/payment transaction.",
        ],
        ui_changes=[
            "Подписка появляется в list и upcoming renewals.",
            "Payment workflow panel показывает planned/task_open/paid/skipped state.",
            "После Mark as paid next billing date продвигается, текущий cycle становится paid.",
            "Action Center item перемещается в resolved/recently updated.",
        ],
        business_rules=[
            "Creating a subscription posts no money.",
            "Attaching a document to subscription posts no money.",
            "subscription-sweep never posts money.",
            "Mark as paid должен быть replay-safe.",
        ],
        edge_cases=[
            "Double click Mark as paid: RPC возвращает existing transaction.",
            "Late payment: next date рассчитывается с сохранением billing anchor.",
            "Subscription cancelled/inactive: новые reminder/payment cycles не должны создаваться.",
            "Duplicate cycle insert: 23505 должен разрешиться в existing reusable cycle.",
            "Legacy quick renew скрывается, когда managed cycle существует, чтобы не продвинуть дату дважды.",
        ],
    ),
    Workflow(
        name="Уведомления, напоминания и read-state",
        purpose="Доставлять пользователю сигналы, не меняя бизнес-состояние простым чтением уведомления.",
        business_value="Пользователь может очистить inbox уведомлений, но обязательства не исчезнут без фактического решения.",
        user_goal="Просмотреть уведомления, отметить прочитанными, настроить push/sound/quiet hours и перейти к нужному объекту.",
        expected_result="notifications.read_at обновлён, но task/payment/subscription/action item остаётся активным до бизнес-действия.",
        modules="Notifications, Reminders, Action Center, Settings",
        related_entities="notifications, reminder_schedules, push_subscriptions, action_items, todos, subscriptions, money_transactions",
        preconditions=[
            "Пользователь вошёл в организацию.",
            "Для push включены browser permission и VAPID config.",
            "Reminder schedules существуют для задач, подписок, planned payments или draft documents.",
        ],
        journey=[
            "Пользователь открывает bell/notification center и видит непрочитанные события.",
            "Пользователь нажимает notification, чтобы перейти на target_url или detail screen.",
            "Пользователь нажимает Mark as read или Mark all as read.",
            "Для настроек пользователь открывает /dashboard/settings/notifications и меняет push, sound, timezone или quiet hours.",
            "Пользователь закрывает само обязательство только через соответствующий workflow: complete task, Mark as paid, resolve/dismiss item.",
        ],
        system_behaviour=[
            "mark_notification_read и mark_all_visible_notifications_read обновляют только notifications.read_at.",
            "Reminder cron claims bounded batch FOR UPDATE SKIP LOCKED и revalidates membership/source state.",
            "Напоминания создают attention item, in-app notification, delivery record и completion state атомарно.",
            "Date changes cancel pending/processing reminder rows and insert a new set; delivered history remains.",
            "Quiet hours suppress disruptive delivery, not durable in-app history.",
        ],
        ui_changes=[
            "Unread badge уменьшается после read action.",
            "Action Center counters не исчезают от read action.",
            "Push/audio могут быть отключены настройками, но in-app история сохраняется.",
            "Ошибки permission/browser permission показываются в notification settings.",
        ],
        business_rules=[
            "Read is not resolved.",
            "Notification is a pointer; obligation is business state.",
            "Snooze hides Action Center item until due time, но не читает и не решает уведомление автоматически.",
            "Reminder idempotency key включает source, recipient, milestone и source date.",
        ],
        edge_cases=[
            "Пользователь нажал Mark all as read: overdue task всё ещё overdue.",
            "Recipient removed/inactive: reminder stop condition предотвращает доставку.",
            "Due date changed: старые pending schedules отменяются.",
            "Cron retry: delivered duplicate должен быть предотвращён idempotency key.",
            "Нет VAPID config: push отключён, in-app уведомления остаются.",
        ],
    ),
    Workflow(
        name="AI Assistant: рекомендации, инсайты и summary",
        purpose="Дать пользователю помощника для анализа данных, но не автономного агента, который меняет бизнес-факты.",
        business_value="AI ускоряет понимание документов, метрик и действий, сохраняя human confirmation для операций.",
        user_goal="Сгенерировать рекомендации или инсайты, прочитать summary, принять или отклонить предложения.",
        expected_result="Появляются insights/recommendations/summaries или planner/document suggestions; бизнес-данные меняются только после user action.",
        modules="AI, Analytics, Documents, Capture Inbox, Action Center, Billing Usage",
        related_entities="ai_insights, ai_recommendations, planner_suggestions, extraction results, action_items",
        preconditions=[
            "ANTHROPIC_API_KEY настроен или включён mock для локальной обработки документов.",
            "Организация имеет доступ к AI suggestions/generation по plan limits.",
            "Данные для анализа существуют в dashboard metrics, documents или captured entries.",
        ],
        journey=[
            "Пользователь открывает /dashboard/ai и запускает generation recommendations или insights.",
            "Система показывает loading state и затем список рекомендаций/инсайтов.",
            "Пользователь читает reasoning и может dismiss recommendation или перейти к связанному workflow.",
            "На document/inbox surfaces AI создаёт reviewable draft или suggestion, а не готовую запись.",
            "Пользователь подтверждает изменения только в целевом модуле: Tasks, Money, Documents, Action Center.",
        ],
        system_behaviour=[
            "AI actions читают агрегированные метрики и создают bounded suggestions/insights.",
            "Provider errors логируются; primary user action не должен молча менять бизнес-состояние.",
            "AI output проходит Zod/schema validation и allow-list маршрутизацию.",
            "Billing feature gate/usage service ограничивает количество AI suggestions.",
            "Suggestions sweep может expire stale pending suggestions.",
        ],
        ui_changes=[
            "AI page показывает loading, empty, success и error states.",
            "Recommendation может исчезнуть или пометиться dismissed после пользовательского решения.",
            "Low-confidence или missing fields отображаются как review/missing information.",
            "Usage limit может показать upgrade/request-access prompt.",
        ],
        business_rules=[
            "AI не постит расходы, не оплачивает подписки и не завершает обязательства автономно.",
            "Never say automatic accounting; корректная формулировка: AI suggests, user confirms.",
            "AI suggestions ниже confidence threshold не должны применяться напрямую.",
            "Paused CRM/Booking не должны входить в активные AI claims.",
        ],
        edge_cases=[
            "Provider timeout: пользователь видит ошибку или pending retry, данные не меняются.",
            "AI вернул невалидный JSON/payload: suggestion rejected/fallback, no mass assignment.",
            "Usage limit exceeded: generation blocked before provider call.",
            "Dismissed recommendation: не должна появляться как новая без новой причины.",
            "Entity deleted before summary: UI должен отказать безопасно.",
        ],
    ),
    Workflow(
        name="Analytics и overview",
        purpose="Показать пользователю сводку по операционной активности, деньгам, задачам, подпискам и событиям.",
        business_value="Помогает владельцу увидеть тренды без ручной сверки всех модулей.",
        user_goal="Открыть overview/analytics, посмотреть метрики, activity timeline и при необходимости создать report/snapshot/widget.",
        expected_result="Дашборд показывает актуальные агрегаты, но не меняет исходные бизнес-объекты без явного действия.",
        modules="Analytics, Dashboard Overview, Domain Events, Money, Tasks, Subscriptions",
        related_entities="domain_events, analytics snapshots, reports, widgets, module records",
        preconditions=[
            "Пользователь открывает /dashboard/overview или /dashboard/analytics.",
            "В организации есть хотя бы часть данных для расчёта метрик.",
            "Для create/update report или widget требуется data.write.",
        ],
        journey=[
            "Пользователь открывает Overview и видит cross-module metrics roll-up.",
            "Пользователь открывает Analytics и выбирает нужный блок: dashboard metrics, module stats, activity timeline.",
            "Если доступна action-кнопка отчёта/snapshot/widget, пользователь нажимает её и задаёт параметры.",
            "Система сохраняет snapshot/report/widget и обновляет соответствующий список.",
            "Пользователь возвращается в Action Center для действий, если аналитика выявила проблему.",
        ],
        system_behaviour=[
            "Analytics queries читают данные других модулей и domain_events.",
            "Snapshots/reports/widgets создаются через server actions с org context.",
            "Dedicated analytics schema ограничена; часть метрик computed at read time.",
            "С ростом данных query cost может увеличиваться, поэтому caching/aggregation остаются future hardening.",
        ],
        ui_changes=[
            "Пустая организация показывает empty states вместо ошибочных графиков.",
            "Activity timeline обновляется после доменных событий.",
            "Созданный snapshot/report появляется в списке или dashboard area.",
            "Ошибки доступа/лимитов отображаются без изменения данных.",
        ],
        business_rules=[
            "Analytics is read-heavy; метрики не являются бухгалтерским первоисточником.",
            "Денежные метрики должны уважать multi-currency и historical FX.",
            "Action Center остаётся primary screen, overview - secondary.",
        ],
        edge_cases=[
            "Недостаточно данных: отображается empty state с понятным next action.",
            "FX incomplete: base totals должны быть помечены incomplete.",
            "Long query/failure: UI должен показать error/loading state.",
            "Snapshot создан во время concurrent updates: snapshot отражает состояние на момент чтения.",
        ],
    ),
    Workflow(
        name="Settings, участники и workspace",
        purpose="Управлять профилем пользователя, настройками рабочей области, участниками, ролями и уведомлениями.",
        business_value="Команда может безопасно распределять доступ и поддерживать актуальные данные организации.",
        user_goal="Обновить профиль/workspace, пригласить участника, изменить роль, удалить участника или настроить уведомления.",
        expected_result="Настройки сохранены, доступы пересчитаны, owner guard и billing limits соблюдены.",
        modules="Settings, Members, Billing, Notifications, Auth",
        related_entities="profiles, organizations, workspaces, memberships, invites, notification_preferences",
        preconditions=[
            "Пользователь открыт на /dashboard/settings или подстраницах profile/workspace/members/notifications/billing.",
            "Profile read/update доступен пользователю для себя.",
            "Workspace, members и billing manage требуют owner/admin или соответствующих permissions.",
        ],
        journey=[
            "Profile: пользователь меняет имя, avatar или персональные настройки и сохраняет.",
            "Workspace: owner/admin меняет название или workspace fields и сохраняет.",
            "Members: owner/admin вводит email, роль и отправляет приглашение.",
            "Member management: owner/admin меняет роль или удаляет участника. Система проверяет owner guard.",
            "Notifications: пользователь меняет preferences и проверяет тестовое уведомление, если доступно.",
        ],
        system_behaviour=[
            "authorizeSettingsAction возвращает context или null без раскрытия чужих данных.",
            "Admin permissions включают workspace.read/update, members.read/invite/update_role/remove, billing.read/manage.",
            "Invite/accept/decline работают через org-scoped RPC и membership checks.",
            "Avatar storage защищён RLS и file validation.",
            "Team member limits могут проверяться через billing usage.",
        ],
        ui_changes=[
            "Формы показывают saved/error states.",
            "Members table обновляет pending invites и active members.",
            "Недоступные actions скрываются или возвращают forbidden state.",
            "Billing/member limit может показать upgrade prompt.",
        ],
        business_rules=[
            "Нельзя удалить последнего владельца или нарушить owner guard.",
            "Members без billing.manage не могут запускать billing actions.",
            "Client payload не доверяется для organization_id.",
            "Invites должны быть tenant-scoped and non-leaking.",
        ],
        edge_cases=[
            "Email уже приглашён: система должна показать duplicate/pending invite state.",
            "Invite token expired/declined: accept не создаёт membership.",
            "Пользователь пытается удалить себя как последнего owner: отказ.",
            "Billing plan seat limit reached: invite запрещён или требует upgrade.",
            "Avatar upload failed: профильные поля не должны частично ломаться.",
        ],
    ),
    Workflow(
        name="Billing, trial, планы и usage limits",
        purpose="Контролировать доступ организации к возможностям платформы, лимитам и платным планам.",
        business_value="SaaS остаётся управляемым: trial, лимиты, paid activation и provider webhooks не смешиваются с бизнес-деньгами пользователя.",
        user_goal="Понять текущий план, лимиты, trial state, перейти на платный план или открыть billing portal, когда это доступно.",
        expected_result="Access state и entitlements отражают локальные billing records; paid activation происходит только через verified Paddle webhook.",
        modules="Billing, Settings Billing, Paddle boundary, Usage Service, Feature Gate",
        related_entities="billing_subscriptions, usage counters, plan catalog, provider events, memberships",
        preconditions=[
            "Пользователь owner/admin с billing.manage открывает /dashboard/settings/billing или /dashboard/settings/plans.",
            "В private_beta checkout и customer portal намеренно отключены.",
            "В paid_beta/production Paddle runtime config обязателен.",
        ],
        journey=[
            "Пользователь открывает billing page и видит текущий plan, usage, trial/access state.",
            "Если доступен upgrade, owner/admin выбирает план Starter, Pro или Business.",
            "Система запускает provider checkout только при paid mode и полной Paddle config.",
            "Success redirect показывает результат, но не активирует plan.",
            "Verified Paddle webhook обновляет local billing subscription; после этого entitlements и usage limits меняются.",
        ],
        system_behaviour=[
            "Local plan catalog является source of truth для limits и pricing copy.",
            "featureGateService и usageService ограничивают documents.process, AI suggestions, storage upload, automations, team seats.",
            "Paddle events проверяются по raw body/signature и idempotent provider_event_id.",
            "Out-of-order provider events не должны откатывать локальное состояние.",
            "Paddle billing events never create Nevora money_transactions.",
        ],
        ui_changes=[
            "Private beta показывает честное сообщение: checkout/portal unavailable.",
            "Usage cards показывают consumed/limit или Unlimited.",
            "Limit denied flow показывает upgrade/request-access prompt.",
            "После webhook activation план и лимиты обновляются.",
        ],
        business_rules=[
            "Checkout success redirect display-only.",
            "Paid activation only through verified Paddle webhook.",
            "Billing organization scope берётся server-side.",
            "CRM и Booking paused, поэтому не участвуют в active monetization claims.",
        ],
        edge_cases=[
            "Missing Paddle config in paid mode: release blocker / controlled error.",
            "Duplicate webhook delivery: no duplicate subscription update.",
            "Invalid signature: reject.",
            "Member without billing.manage: deny before provider call.",
            "Trial expired: write/execute actions могут быть read-only blocked согласно access state.",
        ],
    ),
    Workflow(
        name="Связи между сущностями и reverse navigation",
        purpose="Связать документы, задачи, транзакции и подписки в операционный граф без прямых foreign keys между модулями.",
        business_value="Пользователь видит контекст: какой документ относится к платежу, какая задача связана с подпиской, какая транзакция подтверждена документом.",
        user_goal="Создать или увидеть связь между активными сущностями и перейти к связанному объекту.",
        expected_result="Создан entity_link внутри организации; detail pages показывают related entities без падения при недоступных targets.",
        modules="Relations, Documents, Tasks, Money, Subscriptions, Action Center",
        related_entities="entity_links, documents, todos, money_transactions, subscriptions",
        preconditions=[
            "Source и target относятся к активным entity kinds: task, document, transaction, subscription.",
            "Обе сущности принадлежат текущей организации.",
            "У пользователя есть entity_link.create или соответствующие relation permissions.",
        ],
        journey=[
            "Пользователь открывает detail page документа, задачи, транзакции или подписки.",
            "Пользователь выбирает добавить связь или принимает suggested relation из Action Center/Capture Inbox.",
            "Система ищет/валидирует target entity и показывает label/route.",
            "Пользователь подтверждает связь.",
            "Связанный объект появляется в relation viewer и становится navigable.",
        ],
        system_behaviour=[
            "Relations используют entity_links source_type/source_id -> target_type/target_id.",
            "verifyEntityOrganization fails closed для неизвестных или paused entity types.",
            "RELATION_ENTITY_CONFIG является единым source of truth для table, route и label metadata.",
            "Duplicate active links блокируются partial unique index.",
            "Reverse links на document detail гидратируются и недоступные targets удаляются из выдачи без crash.",
        ],
        ui_changes=[
            "Relation viewer показывает related records, тип связи и ссылку на detail page.",
            "Missing/deleted relation может исчезнуть или отобразиться как unavailable.",
            "Action Center missing_relation item может resolved после создания связи.",
        ],
        business_rules=[
            "CRM/Booking entity kinds не входят в active relation scope.",
            "Links are cross-module context, not ownership transfer.",
            "Удаление target не должно раскрывать данные другого tenant.",
        ],
        edge_cases=[
            "Target принадлежит другой организации: fail closed.",
            "Target удалён после создания link: viewer не падает.",
            "Duplicate link: пользователь получает controlled message.",
            "Paused CRM type: mapping отсутствует, связь запрещена.",
            "Concurrent link creation: unique constraint предотвращает дубль.",
        ],
    ),
    Workflow(
        name="Automation и фоновые repair jobs",
        purpose="Обрабатывать доменные события и запускать фоновые задания, которые чинят или напоминают, но не проводят деньги.",
        business_value="Система остаётся устойчивой к сбоям: застрявшие extraction jobs, reminder schedules и subscription cycles восстанавливаются без ручного обхода.",
        user_goal="Не запускать automation вручную, а видеть корректные action items, reminders и recovered jobs.",
        expected_result="Cron/automation создаёт drafts, reminders или action items; posted money появляется только после user confirmation.",
        modules="Domain Events, Automation, Cron, Documents, Subscriptions, Notifications, Action Center",
        related_entities="domain_events, automation_logs, reminder_schedules, extraction jobs, subscription_payment_cycles, action_items",
        preconditions=[
            "Machine routes настроены и пропущены proxy как MACHINE_ROUTES.",
            "CRON_SECRET установлен; provider webhooks имеют свои подписи.",
            "Domain events валидируются по registered event names и required envelope.",
        ],
        journey=[
            "Пользователь выполняет обычное действие: создаёт документ, подписку, задачу или транзакцию.",
            "После successful primary write модуль emits domain event best-effort.",
            "Automation handler или cron later создаёт/чинит related item: reminder, suggestion, cycle, action item.",
            "Пользователь видит результат в Action Center/Notifications и решает его вручную.",
            "Если background job fail, система логирует ошибку и повторяет/repair при следующем sweep.",
        ],
        system_behaviour=[
            "emitDomainEvent validates envelope и re-checks organizationId equals active org.",
            "domain_events append-only; failed emit logs and returns without breaking primary action.",
            "extraction-sweep recovers stuck extraction jobs; suggestions-sweep expires stale suggestions.",
            "subscription-sweep creates missing cycles/suggestions/action items but never marks paid.",
            "reminders cron creates notifications and attention items with idempotency.",
        ],
        ui_changes=[
            "Пользователь не видит cron, но видит появившиеся reminders/action items.",
            "Failed automation может создать missing_information или error-support context.",
            "Recovered extraction меняет status с stuck/processing на done/failed.",
        ],
        business_rules=[
            "No cron, no AI job, no event handler posts money.",
            "Automation side effects must be idempotent.",
            "Machine routes fail closed without secret/signature.",
            "One event mechanism per table: service emit or DB trigger, not both.",
        ],
        edge_cases=[
            "CRON_SECRET отсутствует: routes return 401/503 and log misconfigured state.",
            "Two workers claim same job: conditional status update/row lock prevents duplicate work.",
            "Handler failure: primary user write remains successful.",
            "Unknown event name: schema rejects emit.",
            "Orphan entity link: future sweep should soft-delete or hide safely.",
        ],
    ),
    Workflow(
        name="Приостановленные модули: CRM и Booking",
        purpose="Документировать, что CRM и Booking существуют в кодовой базе, но не являются активной пользовательской возможностью.",
        business_value="Поддержка и QA не обещают пользователю функции, которые закрыты продуктовым scope и hard-gated.",
        user_goal="Понять, почему пункты CRM/Booking недоступны и что должно происходить при переходе по старым ссылкам.",
        expected_result="Пользователь не получает доступ к paused modules; данные не раскрываются через UI, Server Actions или public APIs.",
        modules="Paused Module Guards, CRM, Booking, Routes, RLS/Grants",
        related_entities="crm entities, booking pages/hosts/services/requests, paused module config",
        preconditions=[
            "Production environment не включает NEVORA_ENABLE_CRM и NEVORA_ENABLE_BOOKING.",
            "Навигация и pricing/landing copy не обещают CRM/Booking.",
            "Тесты paused-modules coverage проверяют новые ungated files.",
        ],
        journey=[
            "Пользователь не видит CRM или Booking в активной навигации.",
            "Если пользователь открывает старую ссылку /dashboard/crm или /dashboard/booking, страница возвращает 404 или controlled unavailable state.",
            "Если внешний посетитель открывает /booking/[organizationSlug], public surface также closed.",
            "Server Actions для CRM/Booking отказывают, даже если page не отрендерилась.",
            "Support объясняет, что модуль paused и не входит в текущий product promise.",
        ],
        system_behaviour=[
            "CRM page gated; CRM Server Actions call assertPausedModuleAction.",
            "Booking dashboard/public pages gated; public API route handlers use pausedModuleGuard.",
            "Migration 098 закрыла Booking anon read/write DB surface and revoked anon EXECUTE on public booking RPCs.",
            "Relations config не мапит CRM entity types; verifyEntityOrganization fails closed.",
            "Reactivation requires explicit product decision and coordinated PR: env flag, nav, pricing, landing, tests, DB grants.",
        ],
        ui_changes=[
            "Пользователь видит 404/unavailable вместо формы или данных.",
            "Нет nav item, pricing entitlement или marketing claim.",
            "Старые persisted target_url не должны раскрыть данные paused module.",
        ],
        business_rules=[
            "Paused means not merely hidden; mutation and data surfaces are closed.",
            "Un-pausing Booking must deliberately restore public DB grants and rate-limit verification.",
            "Paused modules must not appear in active AI, billing or relation claims.",
        ],
        edge_cases=[
            "У организации была опубликована booking page до pause: она больше не обслуживается.",
            "Старый Server Action endpoint вызван напрямую: server-side guard rejects.",
            "Supabase anon key читает booking tables: should get permission denied after 098.",
            "AI summary or Action Center accidentally includes CRM: considered scope leak to clean up.",
            "Env flag включён в production случайно: release checks should catch nav/pricing/test mismatch.",
        ],
    ),
]


def add_module_inventory(doc: Document):
    add_heading(doc, "Карта модулей", 1)
    add_body(doc, "Активная продуктовая модель строится вокруг Action Center. Overview и Analytics помогают понимать состояние, но ежедневные решения принимаются из /dashboard.")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1.35, 2.35, 1.55, 1.25])
    repeat_table_header(table.rows[0])
    headers = ["Модуль", "Пользовательская ценность", "Маршрут", "Статус"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.add_run(header)
        set_paragraph_spacing(p, after=0, line_spacing=1.15)
        for r in p.runs:
            set_run_font(r, size=9.5, color=INK, bold=True)
    for row in MODULE_ROWS:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.add_run(value)
            set_paragraph_spacing(p, after=0, line_spacing=1.1)
            for r in p.runs:
                set_run_font(r, size=9, color=INK)
    doc.add_paragraph()


def add_basics(doc: Document):
    add_heading(doc, "Базовые правила продукта", 1)
    add_callout(doc, "Главное правило денег", "AI предлагает, пользователь подтверждает, модульный сервис выполняет. Ни cron, ни AI job, ни event handler не создают posted money transaction без явного пользовательского подтверждения или утверждённого идемпотентного workflow.")
    add_heading(doc, "Роли и доступ", 2)
    add_bullets(doc, [
        "Owner/admin управляют workspace, участниками, ролями и billing.",
        "Member работает с доступными данными организации, но не получает billing.manage по умолчанию.",
        "Profile read/update доступен самому пользователю.",
        "Server actions получают organization/workspace из server context, а не из формы.",
        "RLS в Supabase является финальной tenant-boundary.",
    ])
    add_heading(doc, "Типовые состояния", 2)
    add_label_table(doc, [
        ("Task", "todo, in_progress, done"),
        ("Financial task", "open, paid, skipped, dismissed"),
        ("Action item", "open, in_progress, snoozed, resolved, dismissed, cancelled, failed"),
        ("Document", "draft, published, archived"),
        ("Money transaction", "posted или planned"),
        ("Planner entry", "captured, processing, suggested, accepted, rejected, archived, failed"),
    ])
    add_heading(doc, "Словарь", 2)
    add_bullets(doc, [
        "Action item - элемент внимания, который ждёт решения или работы пользователя.",
        "Notification - указатель на событие; чтение не решает обязательство.",
        "Planned transaction или planned cycle - прогноз или обязательство, но не факт оплаты.",
        "Posted transaction - фактическая запись в денежном журнале.",
        "Entity link - связь между объектами разных модулей без прямой бизнес-зависимости.",
    ])


def add_workflow(doc: Document, index: int, wf: Workflow):
    add_heading(doc, f"{index}. {wf.name}", 1)
    add_heading(doc, "1. Workflow Overview", 2)
    add_label_table(doc, [
        ("Workflow Name", wf.name),
        ("Purpose", wf.purpose),
        ("Business value", wf.business_value),
        ("User goal", wf.user_goal),
        ("Expected result", wf.expected_result),
        ("Modules involved", wf.modules),
        ("Related entities", wf.related_entities),
    ])

    add_heading(doc, "2. Preconditions", 2)
    add_bullets(doc, wf.preconditions)

    add_heading(doc, "3. Step-by-step User Journey", 2)
    add_numbers(doc, wf.journey)

    add_heading(doc, "4. System Behaviour", 2)
    add_bullets(doc, wf.system_behaviour)

    add_heading(doc, "5. UI Changes", 2)
    add_bullets(doc, wf.ui_changes)

    add_heading(doc, "6. Business Rules", 2)
    add_bullets(doc, wf.business_rules)

    add_heading(doc, "7. Edge Cases", 2)
    add_bullets(doc, wf.edge_cases)


def add_qa_support(doc: Document):
    add_heading(doc, "QA и support-чеклист", 1)
    add_body(doc, "Этот раздел помогает использовать руководство как основу для ручного smoke-теста, support wiki и будущих видео-скриптов.")
    add_heading(doc, "Smoke paths", 2)
    add_numbers(doc, [
        "Register -> onboarding -> /dashboard: пользователь попадает в Action Center и видит empty/initial state.",
        "Create task -> set due date -> observe Action Center due item -> mark done.",
        "Create money account -> create expense -> verify balance, category state and activity timeline.",
        "Upload receipt -> extraction/review -> reject: no money transaction is created.",
        "Upload receipt -> extraction/review -> confirm: exactly one posted transaction is created.",
        "Create subscription -> planned cycle -> Mark as paid double click: exactly one transaction and next cycle.",
        "Mark notification read -> verify Action Center item remains unresolved.",
        "Capture Inbox text -> accept task suggestion -> one task created and suggestion resolved.",
        "Open paused CRM/Booking route -> 404/unavailable, no data leakage.",
    ])
    add_heading(doc, "Support notes", 2)
    add_bullets(doc, [
        "Если пользователь говорит, что уведомление исчезло, проверьте notification read state отдельно от action item state.",
        "Если пользователь говорит, что AI 'создал расход', проверьте review confirmation: extraction должна создать только draft/review item.",
        "Если пользователь видит неправильную сумму в валюте, проверьте exchange_rates и complete flag в base summary.",
        "Если документ завис в processing, проверьте extraction-sweep и status claim history.",
        "Если платный план не активировался, проверьте Paddle webhook, а не success redirect.",
    ])
    add_heading(doc, "Ограничения документа", 2)
    add_bullets(doc, [
        "Руководство отражает состояние репозитория на 12 июля 2026.",
        "Точные подписи кнопок могут отличаться после локализационных правок UI; бизнес-поток и проверки остаются источником истины.",
        "CRM и Booking описаны как paused modules, а не как пользовательские функции.",
        "Перед public launch нужны живые e2e/smoke подтверждения для upload->extract->confirm, reminders, Paddle sandbox и double-click idempotency.",
    ])
    add_heading(doc, "Краткий индекс маршрутов", 2)
    add_label_table(doc, [
        ("/dashboard", "Action Center: основной ежедневный экран."),
        ("/dashboard/overview", "Вторичная сводка метрик."),
        ("/dashboard/inbox", "Capture Inbox и AI-предложения."),
        ("/dashboard/tasks", "Обычные задачи и проекты."),
        ("/dashboard/tasks/financial", "Финансовые задачи и Mark as paid."),
        ("/dashboard/money", "Счета, транзакции, переводы, категории."),
        ("/dashboard/documents", "Документы, загрузки, extraction review."),
        ("/dashboard/subscriptions", "Подписки и циклы платежей."),
        ("/dashboard/ai", "AI Assistant: insights, recommendations, summaries."),
        ("/dashboard/analytics", "Analytics, reports, snapshots, timeline."),
        ("/dashboard/settings", "Profile, workspace, members, billing, notifications, developer access."),
    ], widths=[2.2, 4.3])


def build():
    doc = Document()
    configure_styles(doc)
    setup_header_footer(doc)
    add_cover(doc)
    add_heading(doc, "Как читать это руководство", 1)
    add_body(doc, "Каждый workflow описан одинаково: обзор, предусловия, пошаговый путь пользователя, внутреннее поведение системы, UI-изменения, бизнес-правила и edge cases. Это сделано намеренно: один и тот же текст можно использовать как пользовательскую документацию, QA manual, support guide и основу для onboarding/video scripts.")
    add_body(doc, "Документ не является каталогом страниц. Nevora описана как операционная система малого бизнеса, где главная единица работы - подтверждённое действие пользователя в контексте организации.")
    add_module_inventory(doc)
    add_basics(doc)
    doc.add_page_break()

    add_heading(doc, "Полные пользовательские workflow", 1)
    for i, wf in enumerate(WORKFLOWS, start=1):
        add_workflow(doc, i, wf)
    add_qa_support(doc)

    doc.core_properties.title = "Nevora Business OS - Операционное руководство пользователя"
    doc.core_properties.subject = "Operations Manual, User Guide, QA Manual, Support Documentation"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.comments = "Generated from nevora-sys repository context and product contracts."
    doc.save(OUT)


if __name__ == "__main__":
    build()
